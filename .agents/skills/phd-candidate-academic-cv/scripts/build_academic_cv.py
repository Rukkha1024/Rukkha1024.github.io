#!/usr/bin/env python3
"""Build the private academic CV from a validated, ignored JSON manifest."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_ROOT = Path(__file__).resolve().parents[1]
CONTRACT_PATH = SKILL_ROOT / "references" / "current-cv-contract.md"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(65, 65, 65)
NAVY = RGBColor(0x1F, 0x38, 0x64)
NAVY_HEX = "1F3864"


def contract_version() -> str:
    text = CONTRACT_PATH.read_text(encoding="utf-8")
    match = re.search(r"Contract version:\s*`([^`]+)`", text)
    if not match:
        raise ValueError("current-cv-contract.md has no contract version")
    return match.group(1)


def set_font(run: Any, *, size: float, bold: bool = False, italic: bool = False,
             color: RGBColor = BLACK) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_spacing(paragraph: Any, *, before: float = 0, after: float = 1.5,
                keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.keep_with_next = keep_with_next


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): rel_id})
    run = paragraph._element.makeelement(qn("w:r"))
    rpr = paragraph._element.makeelement(qn("w:rPr"))
    rfonts = paragraph._element.makeelement(
        qn("w:rFonts"), {qn("w:ascii"): "Arial", qn("w:hAnsi"): "Arial"}
    )
    color = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "000000"})
    size = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): "18"})
    rpr.extend([rfonts, color, size])
    text_node = paragraph._element.makeelement(qn("w:t"))
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def add_section_heading(doc: Document, label: str) -> None:
    p = doc.add_paragraph(style="CV Section")
    p.add_run(label.upper())


def add_entry_heading(doc: Document, left: str, right: str) -> None:
    p = doc.add_paragraph(style="CV Entry")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_font(p.add_run(left), size=10.5, bold=True)
    set_font(p.add_run("\t" + right), size=10, color=MUTED)


def add_detail(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph(style="CV Detail")
    set_font(p.add_run(text), size=10, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="CV Bullet")
    set_font(p.add_run(text), size=9.6)


def add_bold_name_paragraph(doc: Document, text: str, name: str, *, url: str | None = None) -> None:
    p = doc.add_paragraph(style="CV Citation")
    cursor = 0
    for match in re.finditer(re.escape(name), text, flags=re.IGNORECASE):
        if match.start() > cursor:
            set_font(p.add_run(text[cursor:match.start()]), size=9.6)
        set_font(p.add_run(text[match.start():match.end()]), size=9.6, bold=True)
        cursor = match.end()
    if cursor < len(text):
        set_font(p.add_run(text[cursor:]), size=9.6)
    if url:
        set_font(p.add_run(" "), size=9.6)
        add_hyperlink(p, "DOI", url)


def configure_document(doc: Document, title: str, version: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    styles = doc.styles
    section_style = styles.add_style("CV Section", WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = normal
    section_style.font.name = "Arial"
    section_style.font.size = Pt(10.5)
    section_style.font.bold = True
    section_style.font.color.rgb = NAVY
    section_style.paragraph_format.space_before = Pt(10)
    section_style.paragraph_format.space_after = Pt(3)
    section_style.paragraph_format.keep_with_next = True
    # 1 pt letter spacing (w:spacing is in twentieths of a point).
    section_rpr = section_style.element.get_or_add_rPr()
    spacing = section_rpr.makeelement(qn("w:spacing"), {qn("w:val"): "20"})
    section_rpr.append(spacing)
    # 0.75 pt navy bottom rule (w:sz is in eighths of a point).
    section_ppr = section_style.element.get_or_add_pPr()
    pbdr = section_ppr.makeelement(qn("w:pBdr"), {})
    bottom = section_ppr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "2",
            qn("w:color"): NAVY_HEX,
        },
    )
    pbdr.append(bottom)
    section_ppr.append(pbdr)

    entry = styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
    entry.base_style = normal
    entry.paragraph_format.space_before = Pt(2.5)
    entry.paragraph_format.space_after = Pt(0)
    entry.paragraph_format.keep_with_next = True

    detail = styles.add_style("CV Detail", WD_STYLE_TYPE.PARAGRAPH)
    detail.base_style = normal
    detail.paragraph_format.space_after = Pt(0.5)
    detail.paragraph_format.keep_with_next = True

    bullet = styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.6)
    bullet.font.color.rgb = BLACK
    bullet.paragraph_format.left_indent = Inches(0.36)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    citation = styles.add_style("CV Citation", WD_STYLE_TYPE.PARAGRAPH)
    citation.base_style = normal
    citation.paragraph_format.left_indent = Inches(0.18)
    citation.paragraph_format.first_line_indent = Inches(-0.18)
    citation.paragraph_format.space_after = Pt(3)

    props = doc.core_properties
    props.title = title
    props.subject = f"Private graduate-application academic CV; contract {version}"
    props.identifier = f"phd-candidate-academic-cv-contract:{version}"
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = ""
    props.keywords = ""


def validate_source(data: dict[str, Any], version: str) -> None:
    if data.get("contract_version") != version:
        raise ValueError(
            f"source contract {data.get('contract_version')!r} does not match {version!r}"
        )
    mode = data.get("mode")
    if mode not in {"draft", "final"}:
        raise ValueError("mode must be draft or final")
    for key in ("name", "contact", "research_profile", "education", "publications",
                "research_experience", "skills", "engagement", "software"):
        if not data.get(key):
            raise ValueError(f"missing required source field: {key}")
    gpa_text = " ".join(str(item.get("gpa", "")) for item in data["education"])
    evidence = data.get("evidence", {})
    if mode == "draft":
        if "provisional" not in gpa_text.lower():
            raise ValueError("draft mode requires a provisional 4.0 GPA label")
    else:
        if "provisional" in gpa_text.lower():
            raise ValueError("final mode forbids provisional GPA wording")
        for key in ("four_point_gpa_result_pdf", "bachelor_degree_certificate_pdf"):
            value = evidence.get(key)
            if not value or not Path(value).is_file():
                raise ValueError(f"final mode requires existing evidence: {key}")


def build(data: dict[str, Any], output: Path) -> None:
    version = contract_version()
    validate_source(data, version)
    doc = Document()
    configure_document(doc, f"{data['name']} Academic CV", version)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=1)
    set_font(p.add_run(data["name"].upper()), size=20, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=4)
    contact = data["contact"]
    items = [contact["location"], contact["email"], contact["website"], contact["github"]]
    set_font(p.add_run("  |  ".join(items)), size=9, color=MUTED)

    add_section_heading(doc, "Research Profile")
    p = doc.add_paragraph()
    set_spacing(p, after=2.5)
    set_font(p.add_run(data["research_profile"]), size=10)

    add_section_heading(doc, "Education")
    for item in data["education"]:
        add_entry_heading(doc, item["institution"], item["dates"])
        add_detail(doc, item["credential"])
        for detail_text in item.get("details", []):
            add_detail(doc, detail_text)
        if item.get("gpa"):
            add_detail(doc, item["gpa"])
    add_detail(doc, data["relevant_coursework"], italic=True)

    add_section_heading(doc, "Peer-Reviewed Journal Articles")
    for publication in data["publications"]:
        add_bold_name_paragraph(
            doc, publication["citation"], data["name"], url=publication.get("url")
        )

    add_section_heading(doc, "Research Experience")
    for item in data["research_experience"]:
        add_entry_heading(doc, item["heading"], item["dates"])
        if item.get("subtitle"):
            add_detail(doc, item["subtitle"], italic=True)
        for bullet_text in item["bullets"]:
            add_bullet(doc, bullet_text)

    add_section_heading(doc, "Technical Skills")
    for index, item in enumerate(data["skills"]):
        p = doc.add_paragraph()
        set_spacing(p, after=1, keep_with_next=index < len(data["skills"]) - 1)
        set_font(p.add_run(item["label"] + ": "), size=9.8, bold=True)
        set_font(p.add_run(item["text"]), size=9.8)

    add_section_heading(doc, "Teaching & Community Engagement")
    for item in data["engagement"]:
        add_entry_heading(doc, item["heading"], item["dates"])
        for bullet_text in item["bullets"]:
            add_bullet(doc, bullet_text)

    add_section_heading(doc, "Selected Research Software")
    for index, item in enumerate(data["software"]):
        p = doc.add_paragraph()
        set_spacing(p, after=1.5, keep_with_next=index < len(data["software"]) - 1)
        set_font(p.add_run(item["name"] + ": "), size=9.6, bold=True)
        add_hyperlink(p, item["display_url"], item["url"])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    build(data, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
