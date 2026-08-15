#!/usr/bin/env python3
"""Audit academic-CV evidence gates, structure, privacy, and rendered PDF parity."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document


SKILL_ROOT = Path(__file__).resolve().parents[1]
CONTRACT_PATH = SKILL_ROOT / "references" / "current-cv-contract.md"
REQUIRED_SECTIONS = [
    "RESEARCH PROFILE",
    "EDUCATION",
    "PEER-REVIEWED JOURNAL ARTICLES",
    "RESEARCH EXPERIENCE",
    "TECHNICAL SKILLS",
    "TEACHING & COMMUNITY ENGAGEMENT",
    "SELECTED RESEARCH SOFTWARE",
]
GENERIC_PROHIBITED = [
    r"date of birth",
    r"student id",
    r"id no\.?",
    r"file no\.?",
    r"certificate no\.?",
    r"issue code",
    r"gender",
]


def current_version() -> str:
    text = CONTRACT_PATH.read_text(encoding="utf-8")
    match = re.search(r"Contract version:\s*`([^`]+)`", text)
    if not match:
        raise ValueError("contract version missing")
    return match.group(1)


def docx_text(doc: Document) -> str:
    values = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def validate_text(text: str, data: dict[str, Any], errors: list[str]) -> None:
    lowered = text.lower()
    for section in REQUIRED_SECTIONS:
        if section not in text:
            errors.append(f"missing section: {section}")
    positions = [text.find(section) for section in REQUIRED_SECTIONS]
    if positions != sorted(positions):
        errors.append("section order does not match contract")
    for pattern in GENERIC_PROHIBITED:
        if re.search(pattern, lowered, flags=re.IGNORECASE):
            errors.append(f"prohibited privacy label found: {pattern}")
    for literal in data.get("privacy", {}).get("prohibited_literals", []):
        if literal and str(literal).lower() in lowered:
            errors.append("a private prohibited literal appears in the artifact")
    if data["mode"] == "draft" and "provisional conversion" not in lowered:
        errors.append("draft artifact lacks provisional GPA label")
    if data["mode"] == "final" and "provisional" in lowered:
        errors.append("final artifact contains provisional wording")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--pdf", type=Path)
    args = parser.parse_args()

    data = json.loads(args.source.read_text(encoding="utf-8"))
    version = current_version()
    errors: list[str] = []
    if data.get("contract_version") != version:
        errors.append("source contract version is stale")

    doc = Document(args.docx)
    expected_id = f"phd-candidate-academic-cv-contract:{version}"
    if doc.core_properties.identifier != expected_id:
        errors.append("DOCX contract identifier is stale or missing")
    if doc.core_properties.author or doc.core_properties.last_modified_by:
        errors.append("DOCX author metadata was not scrubbed")
    if doc.tables:
        errors.append("layout tables are forbidden by the current contract")

    visible = docx_text(doc)
    validate_text(visible, data, errors)

    if args.pdf:
        from pypdf import PdfReader

        reader = PdfReader(args.pdf)
        if len(reader.pages) not in (2, 3):
            errors.append(f"PDF must have 2 or 3 pages, found {len(reader.pages)}")
        pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        validate_text(pdf_text, data, errors)
        for section in REQUIRED_SECTIONS:
            if section not in pdf_text:
                errors.append(f"PDF parity failure for section: {section}")

    result = {
        "ok": not errors,
        "mode": data.get("mode"),
        "contract_version": version,
        "errors": errors,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if not errors else 1


if __name__ == "__main__":
    raise SystemExit(main())
