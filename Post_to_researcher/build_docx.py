"""print_cv.md → Minseok_Cho_CV_print.docx (python-docx). 실행: python3 build_docx.py

print_cv.md의 부분집합만 이해한다: 첫 줄 헤더 div(무시하고 표로 대체), ##, ### **제목** 날짜, 단락,
`- ` 리스트, 4칸 들여쓴 하위 리스트, **bold**, [text](url), <br>(빈 단락), &amp;.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / "print_cv.md"
OUT = HERE / "Minseok_Cho_CV_print.docx"
ASSETS = HERE / "assets"
BLUE = RGBColor(0x11, 0x55, 0xCC)
FONT = "Open Sans"
CONTACT = ["Cheongju, South Korea", "cho9911@gmail.com", "rukkha1024.github.io"]
NAME = "MINSEOK CHO"
INLINE = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")


def set_font(run, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline(par, text: str, size: float, base_bold: bool = False, color: RGBColor | None = None) -> None:
    text = text.replace("&amp;", "&")
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**"):
            set_font(par.add_run(piece[2:-2]), size, True, color)
        elif piece.startswith("["):
            label = re.match(r"\[([^\]]+)\]", piece).group(1)
            set_font(par.add_run(label), size, True, BLUE)
        else:
            set_font(par.add_run(piece), size, base_bold, color)


def tight(par, before: float = 0, after: float = 0, line: float = 1.15) -> None:
    fmt = par.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def bottom_border(par, color: str = "1155CC", size: int = 12) -> None:
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", str(size)), ("w:space", "1"), ("w:color", color)):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def no_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def header(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    left, mid, right = table.rows[0].cells
    for cell, width in ((left, Inches(1.9)), (mid, Inches(4.5)), (right, Inches(1.1))):
        cell.width = width
    # QR codes side by side with captions
    p = left.paragraphs[0]
    tight(p)
    for name in ("qr_site.png", "qr_mail.png"):
        p.add_run().add_picture(str(ASSETS / name), width=Inches(0.8))
        p.add_run("  ")
    cap = left.add_paragraph()
    tight(cap)
    set_font(cap.add_run("  Résumé              Email"), 7.5, color=RGBColor(0x44, 0x44, 0x44))
    # name + contact
    p = mid.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(p, after=4)
    set_font(p.add_run(NAME), 24, True)
    for line in CONTACT:
        q = mid.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tight(q)
        set_font(q.add_run(line), 10)
    # photo
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(p)
    p.add_run().add_picture(str(ASSETS / "profile.jpg"), width=Inches(1.1))
    for cell in (left, mid, right):
        cell.vertical_alignment = 1  # center


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    tight(p, before=8, after=3)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text.replace("&amp;", "&")), 14, True, BLUE)
    bottom_border(p)


def h3(doc: Document, line: str) -> None:
    m = re.match(r"\*\*(.+?)\*\*\s*(.*)", line)
    title, date = (m.group(1), m.group(2)) if m else (line, "")
    p = doc.add_paragraph()
    tight(p, before=4, after=1)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    set_font(p.add_run(title), 11, True)
    if date:
        set_font(p.add_run("\t" + date), 11, True)


def bullet(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    tight(p, after=1)
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    if level == 0:
        set_font(p.add_run("•  "), 10, True)
        add_inline(p, text, 10)
    else:
        p.paragraph_format.keep_with_next = False
        set_font(p.add_run("–  "), 9.5, color=RGBColor(0x33, 0x33, 0x33))
        add_inline(p, text, 9.5, color=RGBColor(0x33, 0x33, 0x33))


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    tight(p, after=1)
    add_inline(p, text, 10)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Inches(0.5))
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    assert lines[0].startswith('<div class="hdr">'), "header div expected on line 1"
    header(doc)
    last_bullet_level: int | None = None
    for raw in lines[1:]:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.strip() == "<br>":
            sp = doc.add_paragraph()
            tight(sp)
            sp.runs or set_font(sp.add_run(""), 4)
            continue
        if line.startswith("## "):
            h2(doc, line[3:])
        elif line.startswith("### "):
            h3(doc, line[4:])
        elif line.startswith("    - "):
            bullet(doc, line[6:], 1)
        elif line.startswith("- "):
            bullet(doc, line[2:], 0)
        else:
            paragraph(doc, line)
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
